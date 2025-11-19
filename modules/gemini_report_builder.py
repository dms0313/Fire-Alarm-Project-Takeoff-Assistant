"""Utilities for rendering Gemini analysis output into a DOCX report."""

from __future__ import annotations

from datetime import datetime
from io import BytesIO
from typing import Any, Dict, Iterable, List

from docx import Document
from docx.shared import Pt


def _format_value(value: Any) -> str:
    if value is None:
        return "Not specified"
    if isinstance(value, (list, tuple, set)):
        filtered = [str(item) for item in value if item]
        return ", ".join(filtered) if filtered else "Not specified"
    return str(value)


def _add_key_value_table(document: Document, rows: Iterable[tuple[str, Any]]) -> None:
    rows = [(key, _format_value(value)) for key, value in rows if key]
    if not rows:
        return

    table = document.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid"
    for (key, value), row in zip(rows, table.rows):
        row.cells[0].text = key
        row.cells[1].text = value


def _add_bullet_list(document: Document, items: Iterable[str]) -> None:
    for item in items:
        if not item:
            continue
        paragraph = document.add_paragraph(item, style="List Bullet")
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_after = Pt(3)


def _format_page_label(page: Any) -> str:
    if page in (None, "", "?"):
        return "Page ?"
    return f"Page {page}"


def _format_label(text: str) -> str:
    if not text:
        return "Detail"
    return text.replace("_", " ").title()


def _has_structured_content(value: Any) -> bool:
    if value is None:
        return False
    if isinstance(value, (list, tuple, set)):
        return any(_has_structured_content(item) for item in value)
    if isinstance(value, dict):
        return any(_has_structured_content(item) for item in value.values())
    if isinstance(value, str):
        return value.strip() != ""
    return True


def _render_structured_value(document: Document, value: Any) -> None:
    if value is None:
        document.add_paragraph("No details supplied.")
        return

    if isinstance(value, (list, tuple, set)):
        if not value:
            document.add_paragraph("No details supplied.")
            return
        for item in value:
            if isinstance(item, dict):
                line = "; ".join(
                    f"{_format_label(key)}: {_format_value(val)}"
                    for key, val in item.items()
                    if val is not None and val != ""
                )
                document.add_paragraph(line or "Details unavailable", style="List Bullet")
            else:
                document.add_paragraph(str(item), style="List Bullet")
        return

    if isinstance(value, dict):
        for key, nested_value in value.items():
            if not _has_structured_content(nested_value):
                continue
            document.add_paragraph(_format_label(key), style="Intense Quote")
            _render_structured_value(document, nested_value)
        return

    document.add_paragraph(str(value))


def _append_structured_summary(document: Document, structured_summary: Dict[str, Any]) -> None:
    if not structured_summary or not isinstance(structured_summary, dict):
        return

    project_details = structured_summary.get("project_details") or {}
    sections = structured_summary.get("sections") or {}
    pitfalls = structured_summary.get("possible_pitfalls") or []

    has_details = _has_structured_content(project_details)
    has_sections = _has_structured_content(sections)
    has_pitfalls = isinstance(pitfalls, list) and any(pitfalls)

    if not any([has_details, has_sections, has_pitfalls]):
        return

    document.add_heading("Fire Alarm Scope Summary", level=1)

    if isinstance(project_details, dict) and has_details:
        _add_key_value_table(
            document,
            [
                ("Project Name", project_details.get("project_name")),
                ("Project Location", project_details.get("project_location")),
                ("Building Type", project_details.get("building_type")),
                ("Construction Type", project_details.get("construction_type")),
                ("Occupancy Groups", project_details.get("occupancy_groups")),
                ("Scope Overview", project_details.get("scope_overview")),
            ],
        )

    if isinstance(sections, dict) and has_sections:
        ordered_sections = [
            ("applicable_codes_standards", "1. Applicable Codes & Standards"),
            ("fire_alarm_equipment_scope", "2. Fire Alarm System Equipment & Scope"),
            ("mechanical_hvac_interface", "3. Mechanical & HVAC Interface"),
            ("elevator_interface", "4. Elevator Interface"),
            ("access_control_door_hardware", "5. Access Control & Door Hardware Interface"),
            ("estimating_notes_inconsistencies", "6. Estimating Notes & Inconsistencies"),
            ("required_modules_summary", "7. Required Modules Summary"),
        ]
        for key, heading in ordered_sections:
            content = sections.get(key)
            if not _has_structured_content(content):
                continue
            document.add_heading(heading, level=2)
            _render_structured_value(document, content)

    if has_pitfalls:
        document.add_heading("Possible Pitfalls / Things to Consider", level=2)
        _add_bullet_list(document, [str(item) for item in pitfalls if item])


def build_gemini_report(results: Dict[str, Any]) -> BytesIO:
    """Render Gemini AI analysis results into a downloadable DOCX file."""

    document = Document()
    heading = document.add_heading("Gemini Fire Alarm Detailed Report", level=0)
    heading.alignment = 0

    analysis_ts = results.get("analysis_timestamp")
    if analysis_ts:
        try:
            parsed = datetime.fromisoformat(analysis_ts)
        except ValueError:
            parsed = None
    else:
        parsed = None

    generated_text = parsed.strftime("%B %d, %Y at %I:%M %p") if parsed else "Unknown"
    intro = document.add_paragraph()
    intro.add_run(f"Generated: {generated_text}\n").bold = True
    total_pages = results.get("total_pages")
    if total_pages is not None:
        intro.add_run(f"Total Pages Reviewed: {total_pages}\n")
    if job_id := results.get("job_id"):
        intro.add_run(f"Gemini Job ID: {job_id}")

    # Project Overview
    project_info = results.get("project_info", {}) or {}
    document.add_heading("Project Overview", level=1)
    _add_key_value_table(
        document,
        [
            ("Project Name", project_info.get("project_name")),
            ("Location", project_info.get("location")),
            ("Project Type", project_info.get("project_type")),
            ("Owner / Client", project_info.get("owner")),
            ("Architect", project_info.get("architect")),
            ("Engineer", project_info.get("engineer")),
            ("Project Number", project_info.get("project_number")),
        ],
    )
    if scope := project_info.get("scope_summary"):
        document.add_paragraph(scope)

    structured_summary = results.get("structured_summary") or {}
    _append_structured_summary(document, structured_summary)

    # Codes & Standards
    document.add_heading("Fire Alarm Codes & Standards", level=1)
    code_requirements = results.get("code_requirements", {}) or {}
    codes = code_requirements.get("fire_alarm_codes") or code_requirements.get("fire_alarm_standards") or []
    if codes:
        _add_bullet_list(document, codes)
    else:
        document.add_paragraph("No fire alarm-specific codes were extracted.")

    # Fire Alarm Focus Pages
    document.add_heading("Fire Alarm Focus Pages", level=1)
    fa_pages: List[int] = results.get("fire_alarm_pages") or []
    if fa_pages:
        document.add_paragraph(
            "Gemini isolated the following sheets as containing electrical or life-safety fire alarm content:"
        )
        _add_bullet_list(document, [f"{_format_page_label(page)}" for page in fa_pages])
    else:
        document.add_paragraph("No specific fire alarm pages were identified.")

    # Fire Alarm Notes
    document.add_heading("Fire Alarm System Notes", level=1)
    fire_alarm_notes = results.get("fire_alarm_notes") or []
    if fire_alarm_notes:
        for note in fire_alarm_notes:
            if not isinstance(note, dict):
                continue
            page = _format_page_label(note.get("page"))
            note_type = note.get("note_type") or "Note"
            content = note.get("content") or "Details unavailable"
            paragraph = document.add_paragraph(style="List Number")
            paragraph.add_run(f"{page} – {note_type}: ").bold = True
            paragraph.add_run(content)
    else:
        document.add_paragraph("No project-specific fire alarm notes were captured.")

    # Mechanical coordination (duct detectors, dampers)
    document.add_heading("Mechanical Coordination (Fire Alarm Tie-Ins)", level=1)
    mechanical_devices = results.get("mechanical_devices") or {}
    mech_sections = [
        ("Duct Detectors", mechanical_devices.get("duct_detectors", [])),
        ("Fire / Smoke Dampers", mechanical_devices.get("dampers", [])),
    ]
    added_mechanical = False
    for title, devices in mech_sections:
        document.add_paragraph(title, style="Intense Quote")
        if devices:
            added_mechanical = True
            for device in devices:
                if not isinstance(device, dict):
                    continue
                page = _format_page_label(device.get("page"))
                parts = [
                    f"{page}",
                    device.get("device_type"),
                    device.get("location"),
                    f"Qty: {device.get('quantity')}" if device.get("quantity") else None,
                    device.get("specifications"),
                ]
                text = " | ".join(filter(None, map(str, parts)))
                document.add_paragraph(text, style="List Bullet")
        else:
            document.add_paragraph("No devices noted in this category.", style="List Bullet")

    if not added_mechanical:
        document.add_paragraph(
            "Gemini did not identify duct detectors or fire/smoke dampers tied into the fire alarm system."
        )

    # Specifications & manufacturers
    document.add_heading("Fire Alarm Specifications", level=1)
    specifications = results.get("specifications") or {}
    if specifications:
        for key, value in specifications.items():
            if not value or key in {"error"}:
                continue
            label = key.replace("_", " ").title()
            paragraph = document.add_paragraph()
            paragraph.add_run(f"{label}: ").bold = True
            paragraph.add_run(_format_value(value))
    else:
        document.add_paragraph("No additional specification details were captured by Gemini.")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

